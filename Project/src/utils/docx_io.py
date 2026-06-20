"""Word (.docx) input/output for the signal report.

- ``extract_headings_from_docx`` reads a skeleton template and returns its heading
  texts in order, so the report structure can be driven by an uploaded template.
- ``render_report_docx`` writes the final report as a Word document by parsing the
  section markdown the report agent already produced (format-agnostic: any report,
  blueprint- or template-driven, renders the same way).
"""

from io import BytesIO
from pathlib import Path
from typing import List, Union

from docx import Document

from src.pipeline.schemas import SignalReport


def extract_headings_from_docx(source: Union[str, Path, bytes, BytesIO]) -> List[str]:
    """Return ordered heading texts from a .docx skeleton.

    A paragraph is treated as a heading when its style name starts with "Heading"
    (Heading 1..9, Title). Empty headings are skipped.
    """
    if isinstance(source, bytes):
        source = BytesIO(source)
    document = Document(source)

    headings: List[str] = []
    for paragraph in document.paragraphs:
        style_name = (paragraph.style.name if paragraph.style else "") or ""
        is_heading = style_name.startswith("Heading") or style_name == "Title"
        text = paragraph.text.strip()
        if is_heading and text:
            headings.append(text)
    return headings


def render_report_docx(signal_report: SignalReport, out_path: Path) -> Path:
    """Render the report's section markdown into a Word document."""
    out_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()

    for raw_line in signal_report.report_markdown.splitlines():
        line = raw_line.rstrip()
        if not line:
            continue

        if line.startswith("# "):
            document.add_heading(line[2:].strip(), level=0)
        elif line.startswith("## "):
            document.add_heading(line[3:].strip(), level=1)
        elif line.startswith("- "):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
        elif line.startswith("**") and line.endswith("**"):
            paragraph = document.add_paragraph()
            run = paragraph.add_run(line.strip("*").strip())
            run.bold = True
        else:
            document.add_paragraph(line)

    document.save(str(out_path))
    return out_path
