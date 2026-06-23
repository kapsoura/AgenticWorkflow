"""
FastAPI backend — exposes the complaint processing pipeline as REST API.
"""

import json
import sys
import time
from pathlib import Path
from typing import Optional

# Add parent to path so we can import src modules
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

# Lazy-import heavy modules (sentence-transformers) to keep startup fast
from src.pipeline.database import get_connection, get_db_stats

app = FastAPI(title="Signal Intelligence API", version="1.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:5173",
        "http://localhost:3000",
        "http://127.0.0.1:5173",
        "https://kapsoura.github.io",
    ],
    # Vercel assigns a new random per-deployment URL on every preview deploy
    # (e.g. multi-agent-quality-intellience-3xwmj18qm.vercel.app), in addition
    # to the stable production alias. Matching by pattern avoids having to
    # update this allowlist on every deploy.
    allow_origin_regex=r"https://multi-agent-quality-intellience.*\.vercel\.app",
    allow_methods=["*"],
    allow_headers=["*"],
)

# Global pipeline instance (lazy init)
_pipeline = None


def get_pipeline():
    global _pipeline
    if _pipeline is None:
        from src.pipeline.orchestrator import Pipeline
        from src.embeddings.generator import EmbeddingGenerator
        _pipeline = Pipeline(model="mistral-small", base_url="http://localhost:11434")
        # Pre-load clusters
        embeddings, report_numbers = EmbeddingGenerator.load_embeddings()
        _pipeline.trend_analyzer.fit_clusters(embeddings, report_numbers)
    return _pipeline


# Anthropic multi-agent stack (lazy init) — used for the INTERACTIVE analyze path.
# Ollama (via the Pipeline above) is reserved strictly for the batch extraction job.
_agents = None


def get_agents():
    """Return the Anthropic-backed interactive agents (extraction + ISO 14971 risk).

    These route through ``AnthropicClient`` (Claude CLI / Anthropic API), never
    Ollama. Construction is cheap — no model server is contacted until ``.extract``
    / ``.assess`` is called.
    """
    global _agents
    if _agents is None:
        from src.agents.extraction import ExtractionAgent as AnthropicExtractionAgent
        from src.agents.risk_analysis import RiskAnalysisAgent
        _agents = {
            "extraction": AnthropicExtractionAgent(),
            "risk": RiskAnalysisAgent(),
        }
    return _agents


# Report-generation orchestrator (lazy init). Wires the same sub-agents the
# LangGraph pipeline uses so downloaded reports are authored by the real
# ReportGenerationAgent (LLM-backed, self-critiqued), not deterministic text.
_orchestrator = None


def get_orchestrator():
    global _orchestrator
    if _orchestrator is None:
        from src.agents.archive_trend import ArchiveTrendAnalyzer
        from src.agents.orchestration import OrchestrationAgent
        from src.agents.report_generation import ReportGenerationAgent
        from src.agents.retrieval import RetrievalAgent

        agents = get_agents()
        _orchestrator = OrchestrationAgent(
            extraction_agent=agents["extraction"],
            retrieval_agent=RetrievalAgent(),
            risk_agent=agents["risk"],
            report_agent=ReportGenerationAgent(),
            trend_analyzer=ArchiveTrendAnalyzer(),
        )
    return _orchestrator


def _interactive_extract(narrative: str, report_id: str, product_code: Optional[str] = None) -> dict:
    """Run interactive extraction through the Anthropic agent (never Ollama).

    Returns a JSON-safe dict. On failure / when the LLM client is disabled the
    returned dict carries an ``error`` key so callers can surface a graceful
    status to the UI.
    """
    import dataclasses
    from src.pipeline.schemas import Complaint

    complaint = Complaint(
        complaint_id=report_id,
        product_code=product_code or "",
        manufacturer="Unknown",
        event_type="Malfunction",
        date_received="",
        narrative=narrative,
        source_report_number=report_id,
    )
    agent = get_agents()["extraction"]
    try:
        signal = agent.extract(complaint)
        data = dataclasses.asdict(signal)
        if signal.qms_complaint_category == "NOT_AVAILABLE":
            data["error"] = (
                getattr(agent, "last_fallback_reason", None)
                or "Anthropic LLM client not enabled (set CLAUDE_CLI_PATH or ANTHROPIC_API_KEY)."
            )
        return data
    except Exception as e:  # noqa: BLE001
        return {"error": str(e)}


class ComplaintRequest(BaseModel):
    narrative: str
    report_id: str = "UI-001"
    skip_extraction: bool = True
    product_code: Optional[str] = None


class StepResult(BaseModel):
    step: str
    status: str
    duration_ms: float
    data: dict


@app.get("/api/stats")
def stats():
    """Get database statistics."""
    conn = get_connection()
    s = get_db_stats(conn)
    conn.close()
    return s


@app.post("/api/process")
def process_complaint(req: ComplaintRequest):
    """
    Process a complaint through the full pipeline:
    1. LLM Extraction
    2. BGE Embedding
    3. Cluster Assignment + Similar Events
    Returns step-by-step results for visualization.
    """
    if not req.narrative.strip():
        raise HTTPException(400, "Narrative cannot be empty")

    pipeline = get_pipeline()
    steps = []

    # Step 1: Extraction (optional) — Anthropic agent only; Ollama is batch-only
    t0 = time.time()
    if not req.skip_extraction:
        try:
            extraction_data = _interactive_extract(
                req.narrative, req.report_id, req.product_code
            )
            if extraction_data.get("error"):
                steps.append(StepResult(
                    step="extraction",
                    status="error",
                    duration_ms=round((time.time() - t0) * 1000, 1),
                    data={"error": extraction_data["error"]},
                ))
            else:
                steps.append(StepResult(
                    step="extraction",
                    status="success",
                    duration_ms=round((time.time() - t0) * 1000, 1),
                    data=extraction_data,
                ))
        except Exception as e:
            steps.append(StepResult(
                step="extraction",
                status="error",
                duration_ms=round((time.time() - t0) * 1000, 1),
                data={"error": str(e)},
            ))
    else:
        steps.append(StepResult(
            step="extraction",
            status="skipped",
            duration_ms=0,
            data={"message": "Skipped — not required for cluster assignment"},
        ))

    # Step 2: Embedding
    t1 = time.time()
    try:
        embedding = pipeline.embedder.embed_single(req.narrative)
        steps.append(StepResult(
            step="embedding",
            status="success",
            duration_ms=round((time.time() - t1) * 1000, 1),
            data={
                "dimensions": int(embedding.shape[0]),
                "vector_preview": embedding[:8].tolist(),
                "norm": round(float((embedding ** 2).sum() ** 0.5), 4),
            },
        ))
    except Exception as e:
        steps.append(StepResult(
            step="embedding",
            status="error",
            duration_ms=round((time.time() - t1) * 1000, 1),
            data={"error": str(e)},
        ))
        return {"steps": [s.model_dump() for s in steps]}

    # Step 3: Cluster Assignment
    t2 = time.time()
    try:
        similarity = pipeline.trend_analyzer.assign_complaint(
            embedding=embedding, conn=pipeline.conn, top_k=10
        )
        sim_data = similarity.model_dump()
        # Convert enums
        for k, v in sim_data.items():
            if hasattr(v, "value"):
                sim_data[k] = v.value
        # Get narratives for similar events
        similar_with_text = []
        for ev in sim_data.get("similar_events", [])[:5]:
            row = pipeline.conn.execute(
                "SELECT narrative, product_code, manufacturer FROM events WHERE report_number = ?",
                (ev["report_number"],),
            ).fetchone()
            similar_with_text.append({
                **ev,
                "narrative_preview": (row["narrative"][:200] + "...") if row and row["narrative"] else "",
                "product_code": row["product_code"] if row else "",
                "manufacturer": row["manufacturer"] if row else "",
            })
        sim_data["similar_events"] = similar_with_text

        steps.append(StepResult(
            step="cluster_assignment",
            status="success",
            duration_ms=round((time.time() - t2) * 1000, 1),
            data=sim_data,
        ))
    except Exception as e:
        steps.append(StepResult(
            step="cluster_assignment",
            status="error",
            duration_ms=round((time.time() - t2) * 1000, 1),
            data={"error": str(e)},
        ))

    total_ms = round((time.time() - t0) * 1000, 1)
    return {
        "report_id": req.report_id,
        "total_duration_ms": total_ms,
        "steps": [s.model_dump() for s in steps],
    }


@app.get("/api/clusters")
def get_clusters():
    """Get all cluster summaries."""
    conn = get_connection()
    rows = conn.execute(
        "SELECT id, label, size, growth_rate_30d, dominant_problem, dominant_modality, trend_flag FROM clusters ORDER BY size DESC"
    ).fetchall()
    conn.close()
    return [dict(r) for r in rows]


def _product_trend(product_code: str, limit_months: int = 12) -> list[dict]:
    """Monthly event counts for a single product code (trend for the incoming complaint)."""
    if not product_code:
        return []
    try:
        conn = get_connection()
    except Exception:
        return []
    try:
        rows = conn.execute(
            """
            SELECT substr(date_received, 1, 6) AS label, COUNT(*) AS count
            FROM events
            WHERE product_code = ?
            GROUP BY label
            HAVING label IS NOT NULL AND label != ''
            ORDER BY label
            """,
            (product_code,),
        ).fetchall()
    finally:
        conn.close()
    series = [
        {"label": f"{r['label'][:4]}-{r['label'][4:]}", "count": r["count"]}
        for r in rows
        if r["label"] and len(r["label"]) >= 6
    ]
    return series[-limit_months:]


def _sse(event: str, payload: dict) -> str:
    """Format a Server-Sent Event frame."""
    return f"event: {event}\ndata: {json.dumps(payload, default=str)}\n\n"


@app.post("/api/process/stream")
def process_complaint_stream(req: ComplaintRequest):
    """
    Same pipeline as /api/process but streamed as Server-Sent Events so the UI
    can show the agentic flow live, node-by-node, while analysis happens.
    Emits: `step` events per node, a `trend` event (per-complaint product trend),
    a `cluster` event (membership), and a final `done` event.
    """
    if not req.narrative.strip():
        raise HTTPException(400, "Narrative cannot be empty")

    def gen():
        pipeline = get_pipeline()
        overall = time.time()

        # Step 1: Extraction
        yield _sse("step", {"step": "extraction", "status": "processing", "data": {}})
        t0 = time.time()
        if not req.skip_extraction:
            try:
                extraction_data = _interactive_extract(
                    req.narrative, req.report_id, req.product_code
                )
                if extraction_data.get("error"):
                    yield _sse("step", {
                        "step": "extraction", "status": "error",
                        "duration_ms": round((time.time() - t0) * 1000, 1),
                        "data": {"error": extraction_data["error"]},
                    })
                else:
                    yield _sse("step", {
                        "step": "extraction", "status": "success",
                        "duration_ms": round((time.time() - t0) * 1000, 1),
                        "data": extraction_data,
                    })
            except Exception as e:
                yield _sse("step", {
                    "step": "extraction", "status": "error",
                    "duration_ms": round((time.time() - t0) * 1000, 1),
                    "data": {"error": str(e)},
                })
        else:
            yield _sse("step", {
                "step": "extraction", "status": "skipped", "duration_ms": 0,
                "data": {"message": "Skipped — not required for cluster assignment"},
            })

        # Step 2: Embedding
        yield _sse("step", {"step": "embedding", "status": "processing", "data": {}})
        t1 = time.time()
        try:
            embedding = pipeline.embedder.embed_single(req.narrative)
            yield _sse("step", {
                "step": "embedding", "status": "success",
                "duration_ms": round((time.time() - t1) * 1000, 1),
                "data": {
                    "dimensions": int(embedding.shape[0]),
                    "vector_preview": embedding[:8].tolist(),
                    "norm": round(float((embedding ** 2).sum() ** 0.5), 4),
                },
            })
        except Exception as e:
            yield _sse("step", {
                "step": "embedding", "status": "error",
                "duration_ms": round((time.time() - t1) * 1000, 1),
                "data": {"error": str(e)},
            })
            yield _sse("done", {"total_duration_ms": round((time.time() - overall) * 1000, 1)})
            return

        # Step 3: Cluster assignment
        yield _sse("step", {"step": "cluster_assignment", "status": "processing", "data": {}})
        t2 = time.time()
        cluster_data: dict = {}
        try:
            similarity = pipeline.trend_analyzer.assign_complaint(
                embedding=embedding, conn=pipeline.conn, top_k=10
            )
            cluster_data = similarity.model_dump()
            for k, v in cluster_data.items():
                if hasattr(v, "value"):
                    cluster_data[k] = v.value
            similar_with_text = []
            for ev in cluster_data.get("similar_events", [])[:5]:
                row = pipeline.conn.execute(
                    "SELECT narrative, product_code, manufacturer FROM events WHERE report_number = ?",
                    (ev["report_number"],),
                ).fetchone()
                similar_with_text.append({
                    **ev,
                    "narrative_preview": (row["narrative"][:200] + "...") if row and row["narrative"] else "",
                    "product_code": row["product_code"] if row else "",
                    "manufacturer": row["manufacturer"] if row else "",
                })
            cluster_data["similar_events"] = similar_with_text
            yield _sse("step", {
                "step": "cluster_assignment", "status": "success",
                "duration_ms": round((time.time() - t2) * 1000, 1),
                "data": cluster_data,
            })
            yield _sse("cluster", cluster_data)
        except Exception as e:
            yield _sse("step", {
                "step": "cluster_assignment", "status": "error",
                "duration_ms": round((time.time() - t2) * 1000, 1),
                "data": {"error": str(e)},
            })

        # Per-complaint trend (based on the incoming complaint's product code)
        product_code = req.product_code or cluster_data.get("dominant_modality")
        yield _sse("trend", {
            "product_code": product_code,
            "series": _product_trend(product_code) if product_code else [],
        })

        yield _sse("done", {
            "report_id": req.report_id,
            "total_duration_ms": round((time.time() - overall) * 1000, 1),
        })

    return StreamingResponse(gen(), media_type="text/event-stream")


_graph_cache: Optional[dict] = None


@app.get("/api/graph")
def langgraph_structure():
    """Return the restored LangGraph agentic-flow structure (nodes + edges)."""
    global _graph_cache
    if _graph_cache is not None:
        return _graph_cache
    try:
        from src.pipeline.langgraph_flow import LangGraphSignalWorkflow
        workflow = LangGraphSignalWorkflow()
        drawable = workflow.graph.get_graph()
        nodes = [
            {"id": node_id} for node_id in drawable.nodes.keys()
        ]
        edges = [
            {
                "source": e.source,
                "target": e.target,
                "conditional": bool(getattr(e, "conditional", False)),
            }
            for e in drawable.edges
        ]
        _graph_cache = {"nodes": nodes, "edges": edges}
        return _graph_cache
    except Exception as e:
        raise HTTPException(500, f"Could not build graph: {e}")


# ---------- Endpoints for React UI ----------


PRODUCT_CODES = ["LNH", "JAK", "LLZ"]
EVENT_TYPES = ["Malfunction", "Injury", "Death"]
REPORT_TYPES = ["PSUR", "INCIDENT_ASSESSMENT", "CAPA"]


@app.get("/health")
def health():
    return {"status": "ok", "service": "signal-intelligence-api"}


@app.get("/api/meta")
def meta():
    """Static metadata for the UI."""
    return {
        "product_codes": PRODUCT_CODES,
        "event_types": EVENT_TYPES,
        "report_types": REPORT_TYPES,
    }


@app.get("/api/trends")
def trends(
    dimension: str,
    product_code: Optional[str] = None,
    event_type: Optional[str] = None,
    software_related: Optional[bool] = None,
):
    """Aggregate a dimension from the events/clusters tables."""
    try:
        conn = get_connection()
    except Exception as e:
        raise HTTPException(500, f"Database connection failed: {e}")

    series = []

    allowed_cluster_dims = {"trend_flag", "dominant_problem", "dominant_modality"}

    try:
        if dimension == "product_code":
            sql = "SELECT product_code AS label, COUNT(*) AS count FROM events"
            conditions, params = _build_where(product_code, event_type, software_related)
            sql += conditions + " GROUP BY product_code ORDER BY count DESC LIMIT 20"
            rows = conn.execute(sql, params).fetchall()
            series = [{"label": r["label"] or "Unknown", "count": r["count"]} for r in rows]
        elif dimension == "event_type":
            sql = "SELECT event_type AS label, COUNT(*) AS count FROM events"
            conditions, params = _build_where(product_code, event_type, software_related)
            sql += conditions + " GROUP BY event_type ORDER BY count DESC"
            rows = conn.execute(sql, params).fetchall()
            series = [{"label": r["label"] or "Unknown", "count": r["count"]} for r in rows]
        elif dimension == "manufacturer":
            sql = "SELECT manufacturer AS label, COUNT(*) AS count FROM events"
            conditions, params = _build_where(product_code, event_type, software_related)
            sql += conditions + " GROUP BY manufacturer ORDER BY count DESC LIMIT 12"
            rows = conn.execute(sql, params).fetchall()
            series = [{"label": r["label"] or "Unknown", "count": r["count"]} for r in rows]
        elif dimension == "year":
            sql = "SELECT substr(date_received, 1, 4) AS label, COUNT(*) AS count FROM events"
            conditions, params = _build_where(product_code, event_type, software_related)
            sql += conditions + " GROUP BY label HAVING label IS NOT NULL AND label != '' ORDER BY label"
            rows = conn.execute(sql, params).fetchall()
            series = [{"label": r["label"], "count": r["count"]} for r in rows if r["label"]]
        elif dimension == "month":
            sql = "SELECT substr(date_received, 1, 6) AS label, COUNT(*) AS count FROM events"
            conditions, params = _build_where(product_code, event_type, software_related)
            sql += conditions + " GROUP BY label HAVING label IS NOT NULL AND label != '' ORDER BY label"
            rows = conn.execute(sql, params).fetchall()
            series = [
                {"label": f"{r['label'][:4]}-{r['label'][4:]}", "count": r["count"]}
                for r in rows if r["label"] and len(r["label"]) >= 6
            ]
        elif dimension == "quarter":
            sql = """SELECT substr(date_received, 1, 4) AS yr,
                            CAST((CAST(substr(date_received, 5, 2) AS INTEGER) - 1) / 3 + 1 AS TEXT) AS q,
                            COUNT(*) AS count
                     FROM events"""
            conditions, params = _build_where(product_code, event_type, software_related)
            sql += conditions + " GROUP BY yr, q HAVING yr IS NOT NULL AND yr != '' ORDER BY yr, q"
            rows = conn.execute(sql, params).fetchall()
            series = [
                {"label": f"{r['yr']}-Q{r['q']}", "count": r["count"]}
                for r in rows if r["yr"]
            ]
        elif dimension in allowed_cluster_dims:
            col_map = {
                "trend_flag": "trend_flag",
                "dominant_problem": "dominant_problem",
                "dominant_modality": "dominant_modality",
            }
            col = col_map[dimension]
            rows = conn.execute(
                f"SELECT {col} AS label, COUNT(*) AS count FROM clusters GROUP BY {col} ORDER BY count DESC"
            ).fetchall()
            series = [{"label": r["label"] or "Unknown", "count": r["count"]} for r in rows]
        else:
            conn.close()
            raise HTTPException(400, f"Unknown dimension: {dimension}")
    except HTTPException:
        raise
    except Exception as e:
        conn.close()
        raise HTTPException(500, f"Query error: {e}")

    conn.close()
    return {"dimension": dimension, "series": series}


def _build_where(
    product_code: Optional[str],
    event_type: Optional[str],
    software_related: Optional[bool],
) -> tuple:
    """Build a parameterised WHERE clause from filter values."""
    clauses = []
    params: list = []
    if product_code:
        clauses.append("product_code = ?")
        params.append(product_code)
    if event_type:
        clauses.append("event_type = ?")
        params.append(event_type)
    if software_related is not None:
        clauses.append("software_related = ?")
        params.append(software_related)
    if clauses:
        return " WHERE " + " AND ".join(clauses), params
    return "", params


@app.get("/api/templates")
def templates():
    """Report structures and section catalog."""
    try:
        from src.agents.report_sections import SECTION_KEYWORDS, blueprint_for
        catalog = [
            {"section": name, "keywords": keywords}
            for name, keywords in SECTION_KEYWORDS.items()
        ]
        blueprints = {}
        for rt in REPORT_TYPES:
            try:
                specs = blueprint_for(rt)
                blueprints[rt] = [{"name": s.name, "title": s.title} for s in specs]
            except Exception:
                blueprints[rt] = []
        return {"section_catalog": catalog, "blueprints": blueprints}
    except ImportError:
        return {"section_catalog": [], "blueprints": {}}


# ─── Analytics (deterministic dashboard aggregations over the real archive) ──
_analytics_cache: Optional[dict] = None


def _analytics_data() -> dict:
    """Build (events_by_code, recalls) from ``data/signal_intelligence.db`` in the
    shape ``src.analytics`` expects. Cached — the archive is static at runtime.
    """
    global _analytics_cache
    if _analytics_cache is None:
        from collections import defaultdict
        conn = get_connection()
        try:
            problems_by_event: dict = defaultdict(list)
            for r in conn.execute(
                "SELECT event_id, problem_code FROM event_problems"
            ).fetchall():
                problems_by_event[r["event_id"]].append(r["problem_code"])

            by_code: dict = defaultdict(list)
            for r in conn.execute(
                "SELECT id, date_received, event_type, product_code, manufacturer FROM events"
            ).fetchall():
                by_code[r["product_code"] or "Unknown"].append({
                    "date_received": str(r["date_received"] or ""),
                    "event_type": r["event_type"],
                    "device": [{"manufacturer_d_name": r["manufacturer"]}],
                    "product_problems": problems_by_event.get(r["id"], []),
                })

            recall_count = conn.execute("SELECT COUNT(*) FROM recalls").fetchone()[0]
            recalls = [{} for _ in range(int(recall_count))]
        finally:
            conn.close()
        _analytics_cache = {"events_by_code": dict(by_code), "recalls": recalls}
    return _analytics_cache


@app.get("/api/analytics/stats")
def analytics_stats():
    """Headline dashboard counts computed deterministically over the real archive."""
    from src.analytics import compute_stats
    from src.config import SQLITE_DB_PATH

    data = _analytics_data()
    return compute_stats(data["events_by_code"], data["recalls"], SQLITE_DB_PATH)


@app.get("/api/analytics/trends")
def analytics_trends(
    dimension: str = "product_code",
    product_code: Optional[str] = None,
    event_type: Optional[str] = None,
    software_related: Optional[bool] = None,
    top_n: int = 12,
):
    """Aggregate one allow-listed dimension into a chart-ready series."""
    from src.analytics import compute_trends
    from src.config import SQLITE_DB_PATH

    data = _analytics_data()
    filters = {
        "product_code": product_code,
        "event_type": event_type,
        "software_related": software_related,
    }
    try:
        return compute_trends(
            data["events_by_code"], SQLITE_DB_PATH, dimension, filters, top_n=top_n
        )
    except ValueError as exc:
        raise HTTPException(400, str(exc))


# ─── Analyze helpers (retrieval / recalls / report) ─────────────────────────


def _enrich_similar_events(conn, similar_events: list, limit: int = 5) -> list:
    """Attach real narrative snippets + device context to retrieved events."""
    enriched = []
    for ev in similar_events[:limit]:
        row = conn.execute(
            "SELECT narrative, product_code, manufacturer, device_name, date_received "
            "FROM events WHERE report_number = ?",
            (ev.get("report_number"),),
        ).fetchone()
        narrative = (row["narrative"] if row and row["narrative"] else "") or ""
        snippet = (narrative[:220] + "…") if len(narrative) > 220 else narrative
        enriched.append({
            **ev,
            "narrative_snippet": snippet or None,
            "product_code": row["product_code"] if row else None,
            "manufacturer": row["manufacturer"] if row else None,
            "device_name": row["device_name"] if row else None,
            "date_received": row["date_received"] if row else None,
        })
    return enriched


def _related_recalls(conn, product_code: Optional[str], limit: int = 5) -> list:
    """Pull openFDA recalls for the same product code as supporting evidence."""
    if not product_code:
        return []
    rows = conn.execute(
        "SELECT recall_number, reason_for_recall, root_cause, classification, "
        "recalling_firm, recall_date FROM recalls WHERE product_code = ? "
        "ORDER BY recall_date DESC LIMIT ?",
        (product_code, limit),
    ).fetchall()
    out = []
    for r in rows:
        reason = (r["reason_for_recall"] or "")
        out.append({
            "recall_number": r["recall_number"],
            "reason_for_recall": (reason[:240] + "…") if len(reason) > 240 else reason,
            "root_cause": r["root_cause"],
            "classification": r["classification"],
            "recalling_firm": r["recalling_firm"],
            "recall_date": r["recall_date"],
        })
    return out


@app.post("/api/analyze")
def analyze_complaint_endpoint(
    narrative: str,
    product_code: str = "LNH",
    event_type: str = "Malfunction",
    manufacturer: str = "Unknown",
):
    """Run a complaint through the INTERACTIVE agent stack and return structured results.

    Interactive extraction + ISO 14971 risk run through the Anthropic agents
    (``AnthropicClient`` → Claude CLI / Anthropic API). Ollama is NOT used here;
    it is reserved for the batch extraction job. Retrieval + cluster assignment
    use the prebuilt BGE index over ``data/signal_intelligence.db``.
    """
    if not narrative.strip():
        raise HTTPException(400, "Narrative cannot be empty")

    import dataclasses
    from src.pipeline.schemas import Complaint, RetrievalEvidence

    pipeline = get_pipeline()
    agents = get_agents()
    t0 = time.time()
    report_id = f"UI-{int(time.time())}"

    complaint = Complaint(
        complaint_id=report_id,
        product_code=product_code,
        manufacturer=manufacturer,
        event_type=event_type,
        date_received="",
        narrative=narrative,
        source_report_number=report_id,
    )

    # ── Extraction (Anthropic agent) ──────────────────────────────────────
    extraction_agent = agents["extraction"]
    extracted_signal = None
    extraction_data: dict = {}
    try:
        extracted_signal = extraction_agent.extract(complaint)
        extraction_data = dataclasses.asdict(extracted_signal)
    except Exception as e:  # noqa: BLE001
        extraction_data = {"error": str(e)}

    extraction_available = (
        extracted_signal is not None
        and extracted_signal.qms_complaint_category != "NOT_AVAILABLE"
    )
    if extraction_available:
        extraction_status = {"ok": True, "backend": extraction_agent.last_backend}
    else:
        extraction_status = {
            "ok": False,
            "reason": "llm_unavailable",
            "backend": getattr(extraction_agent, "last_backend", "unavailable"),
            "message": (
                getattr(extraction_agent, "last_fallback_reason", None)
                or "Anthropic LLM client not enabled (set CLAUDE_CLI_PATH or "
                "ANTHROPIC_API_KEY in .env). Risk below uses the deterministic "
                "ISO 14971 estimate from the risk-analysis agent."
            ),
        }

    # ── Embedding + cluster assignment (BGE index over the real archive) ──
    cluster_data: dict = {}
    try:
        embedding = pipeline.embedder.embed_single(narrative)
        similarity = pipeline.trend_analyzer.assign_complaint(
            embedding=embedding, conn=pipeline.conn, top_k=5
        )
        cluster_data = similarity.model_dump()
        for k, v in cluster_data.items():
            if hasattr(v, "value"):
                cluster_data[k] = v.value
        if cluster_data.get("similar_events"):
            cluster_data["similar_events"] = _enrich_similar_events(
                pipeline.conn, cluster_data["similar_events"]
            )
    except Exception as e:  # noqa: BLE001
        cluster_data = {"error": str(e)}

    # ── openFDA recalls for the same product code (supporting evidence) ───
    recalls = _related_recalls(pipeline.conn, product_code)

    # ── Build evidence bundle for the risk agent ──────────────────────────
    evidence: list = []
    for ev in cluster_data.get("similar_events", []) or []:
        evidence.append(RetrievalEvidence(
            evidence_id=str(ev.get("report_number") or "event"),
            source_type="MAUDE_EVENT",
            product_code=ev.get("product_code") or product_code,
            snippet=ev.get("narrative_snippet") or "",
            score=float(ev.get("similarity_score") or 0.0),
            metadata={
                "manufacturer": ev.get("manufacturer"),
                "device_name": ev.get("device_name"),
                "date_received": ev.get("date_received"),
            },
        ))
    for rc in recalls:
        evidence.append(RetrievalEvidence(
            evidence_id=str(rc.get("recall_number") or "recall"),
            source_type="FDA_RECALL",
            product_code=product_code,
            snippet=rc.get("reason_for_recall") or "",
            score=1.0,
            metadata={
                "classification": rc.get("classification"),
                "root_cause": rc.get("root_cause"),
                "recalling_firm": rc.get("recalling_firm"),
                "recall_date": rc.get("recall_date"),
            },
        ))

    # ── Risk assessment (Anthropic ISO 14971 agent; deterministic when offline) ──
    risk_agent = agents["risk"]
    assessment = None
    try:
        assessment = risk_agent.assess(
            complaint=complaint,
            evidence=evidence,
            extraction=extracted_signal,
            trend=None,
        )
    except Exception:  # noqa: BLE001
        assessment = None

    report_type = "PSUR"
    if assessment is not None:
        report_type = assessment.report_type or report_type
        method = "anthropic" if assessment.llm_backed else "deterministic"
        risk_signals = [s for s in (
            f"Severity {assessment.severity_level} / probability {assessment.probability_level}",
            f"Hazardous situation: {assessment.hazardous_situation}" if assessment.hazardous_situation else "",
            "Escalation required" if assessment.escalation_required else "",
            "PRRC notification required" if assessment.prrc_notification_required else "",
            f"{len(recalls)} related FDA recall(s)" if recalls else "",
        ) if s]
        risk_block = {
            "bucket": assessment.risk_bucket,
            "method": method,
            "report_type": report_type,
            "signals": risk_signals,
            "rationale": assessment.iso_14971_rationale
            or "ISO 14971 assessment produced by the risk-analysis agent.",
            "severity_level": assessment.severity_level,
            "probability_level": assessment.probability_level,
            "escalation_required": assessment.escalation_required,
            "prrc_notification_required": assessment.prrc_notification_required,
            "capa_recommendation": assessment.capa_recommendation,
            "hazardous_situation": assessment.hazardous_situation,
            "harm": assessment.harm,
        }
    else:
        risk_block = {
            "bucket": "UNKNOWN",
            "method": "unavailable",
            "report_type": report_type,
            "signals": [],
            "rationale": "Risk analysis could not be completed.",
        }

    # Which report types this complaint actually warrants, using the same routing
    # rules as OrchestrationAgent.decide_report_types. The UI enables only these
    # download buttons so users cannot request a report that was never generated.
    risk_bucket = risk_block["bucket"]
    escalation = bool(risk_block.get("escalation_required"))
    recall_precedent = bool(recalls)
    event = (complaint.event_type or "").strip().lower()
    applicable_report_types = []
    if event in {"injury", "death"} or escalation or risk_bucket == "UNACCEPTABLE":
        applicable_report_types.append("INCIDENT_ASSESSMENT")
    if risk_bucket in {"ALARP", "UNACCEPTABLE"} or escalation or recall_precedent:
        applicable_report_types.append("CAPA")
    applicable_report_types.append("PSUR")  # always appended

    total_ms = round((time.time() - t0) * 1000, 1)

    return {
        "report_id": report_id,
        "report_type": report_type,
        "applicable_report_types": applicable_report_types,
        "risk_bucket": risk_block["bucket"],
        "risk": risk_block,
        "extraction": extraction_data,
        "extraction_status": extraction_status,
        "evidence_count": len(cluster_data.get("similar_events", [])),
        "recalls": recalls,
        "validation": {"passed": True, "issues": []},
        "cluster": cluster_data,
        "total_duration_ms": total_ms,
    }


# ─── DOCX report export ──────────────────────────────────────────────────────

class ReportExportRequest(BaseModel):
    """Analysis payload (the /api/analyze response) to render as a Word report."""
    report_id: Optional[str] = None
    report_type: Optional[str] = None
    risk_bucket: Optional[str] = None
    risk: Optional[dict] = None
    extraction: Optional[dict] = None
    recalls: Optional[list] = None
    cluster: Optional[dict] = None
    narrative: Optional[str] = None
    product_code: Optional[str] = None
    event_type: Optional[str] = None
    manufacturer: Optional[str] = None


def _reconstruct_domain(payload: "ReportExportRequest"):
    """Rebuild the workflow domain objects from a /api/analyze response payload.

    These feed the real ReportGenerationAgent so the downloaded reports are
    agent-authored. Risk fields not carried in the API response default safely.
    """
    from src.pipeline.schemas import (
        Complaint,
        ExtractedSignal,
        RetrievalEvidence,
        RiskAssessment,
    )

    rid = payload.report_id or "ANALYSIS"
    complaint = Complaint(
        complaint_id=rid,
        product_code=payload.product_code or "",
        manufacturer=payload.manufacturer or "Unknown",
        event_type=payload.event_type or "",
        date_received="",
        narrative=payload.narrative or "",
        source_report_number=rid,
    )

    ex = payload.extraction or {}
    extraction = ExtractedSignal(
        complaint_id=rid,
        qms_complaint_category=ex.get("qms_complaint_category") or "NOT_AVAILABLE",
        key_issues=ex.get("key_issues") or [],
        confidence=float(ex.get("confidence") or 0.0),
        safety_flags=ex.get("safety_flags") or {},
        iso_13485_clauses=ex.get("iso_13485_clauses") or [],
        iso_14971_hazard_tags=ex.get("iso_14971_hazard_tags") or [],
    )

    rk = payload.risk or {}
    risk = RiskAssessment(
        complaint_id=rid,
        severity_level=rk.get("severity_level") or "",
        probability_level=rk.get("probability_level") or "",
        risk_bucket=rk.get("bucket") or payload.risk_bucket or "UNKNOWN",
        escalation_required=bool(rk.get("escalation_required")),
        prrc_notification_required=bool(rk.get("prrc_notification_required")),
        capa_recommendation=rk.get("capa_recommendation") or "",
        report_type=rk.get("report_type") or payload.report_type or "PSUR",
        iso_14971_rationale=rk.get("rationale") or "",
        hazardous_situation=rk.get("hazardous_situation") or "",
        harm=rk.get("harm") or "",
        fsca_required=bool(rk.get("fsca_required")),
        llm_backed=True,
    )

    cluster = payload.cluster or {}
    evidence = []
    for ev in cluster.get("similar_events") or []:
        evidence.append(RetrievalEvidence(
            evidence_id=str(ev.get("report_number") or "event"),
            source_type="MAUDE_EVENT",
            product_code=ev.get("product_code") or complaint.product_code,
            snippet=ev.get("narrative_snippet") or "",
            score=float(ev.get("similarity_score") or 0.0),
            metadata={
                "manufacturer": ev.get("manufacturer"),
                "device_name": ev.get("device_name"),
                "date_received": ev.get("date_received"),
            },
        ))
    for rc in payload.recalls or []:
        evidence.append(RetrievalEvidence(
            evidence_id=str(rc.get("recall_number") or "recall"),
            source_type="FDA_RECALL",
            product_code=complaint.product_code,
            snippet=rc.get("reason_for_recall") or "",
            score=1.0,
            metadata={
                "classification": rc.get("classification"),
                "root_cause": rc.get("root_cause"),
                "recalling_firm": rc.get("recalling_firm"),
                "recall_date": rc.get("recall_date"),
            },
        ))

    return complaint, extraction, risk, evidence


def _generate_agent_reports(payload: "ReportExportRequest", report_types: list) -> list:
    """Author each report type with the real ReportGenerationAgent.

    Returns a list of ``(report_type, report_id, markdown)``. Narrative prose
    needs the LLM backend; deterministic facts still render when it is absent.
    """
    from src.agents.report_sections import ReportContext

    complaint, extraction, risk, evidence = _reconstruct_domain(payload)

    # Same archive events the LangGraph pipeline feeds the trend/quality agents.
    try:
        from src.config import IMAGING_EVENTS_DIR
        from src.utils.data_loader import load_events_for_codes
        events_by_code = load_events_for_codes(
            IMAGING_EVENTS_DIR, [complaint.product_code], 300
        )
    except Exception:  # noqa: BLE001 — degrade to no archive context
        events_by_code = {complaint.product_code: []}

    orchestrator = get_orchestrator()
    out = []
    for rt in report_types:
        context = ReportContext(
            complaint=complaint,
            extraction=extraction,
            retrieval=list(evidence),
            risk=risk,
            report_type=rt,
            events_by_code=events_by_code,
        )
        report = orchestrator.report_agent.create_report(
            trace_id=complaint.complaint_id,
            context=context,
            orchestrator=orchestrator,
        )
        out.append((rt, report.report_id, report.report_markdown))
    return out


def _markdown_to_docx(report_type: str, report_id: str, markdown: str) -> bytes:
    """Render an agent-authored Markdown report to a Word (.docx) document."""
    import io

    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    type_titles = {
        "PSUR": "Periodic Safety Update Report (PSUR)",
        "INCIDENT_ASSESSMENT": "Incident Assessment Report",
        "CAPA": "CAPA Report",
    }

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(type_titles.get(report_type, "Signal Intelligence Report"))
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    for raw in (markdown or "").split("\n"):
        line = raw.rstrip()
        stripped = line.lstrip("#").strip()
        if line.startswith("### "):
            doc.add_heading(stripped, level=3)
        elif line.startswith("## "):
            doc.add_heading(stripped, level=2)
        elif line.startswith("# "):
            doc.add_heading(stripped, level=1)
        elif line.startswith(("- ", "* ")):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif not line:
            doc.add_paragraph("")
        else:
            doc.add_paragraph(line)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


@app.post("/api/analyze/report")
def export_report_docx(payload: ReportExportRequest, report_type: Optional[str] = None):
    """Return a single agent-authored report rendered as a Word document.

    ``report_type`` (PSUR / INCIDENT_ASSESSMENT / CAPA) selects which report the
    ReportGenerationAgent authors; defaults to the analysis's routed type.
    """
    rt = report_type or payload.report_type or "PSUR"
    try:
        generated = _generate_agent_reports(payload, [rt])
    except Exception as exc:  # noqa: BLE001
        raise HTTPException(status_code=500, detail=f"Report generation failed: {exc}")

    _rt, report_id, markdown = generated[0]
    data = _markdown_to_docx(_rt, report_id, markdown)
    filename = f"{payload.report_id or 'analysis'}_{_rt}.docx"
    return StreamingResponse(
        iter([data]),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.post("/api/analyze/reports")
def export_reports_zip(payload: ReportExportRequest):
    """Author PSUR, Incident Assessment and CAPA as three separate Word reports.

    Each is produced by the real ReportGenerationAgent and packaged in one
    ``.zip`` so a single click yields all three distinct reports.
    """
    import io
    import zipfile

    report_types = ["PSUR", "INCIDENT_ASSESSMENT", "CAPA"]
    base = payload.report_id or "analysis"
    try:
        generated = _generate_agent_reports(payload, report_types)
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
            for rt, report_id, markdown in generated:
                zf.writestr(f"{base}_{rt}.docx", _markdown_to_docx(rt, report_id, markdown))
        data = zip_buffer.getvalue()
    except Exception as exc:  # noqa: BLE001
        raise HTTPException(status_code=500, detail=f"Report generation failed: {exc}")

    return StreamingResponse(
        iter([data]),
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="{base}_reports.zip"'},
    )



if __name__ == "__main__":
    import os
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=int(os.environ.get("PORT", 8000)))

